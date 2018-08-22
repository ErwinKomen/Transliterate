# ----------------------------------------------------------------------------------
# Name :    academic2phonemic
# Goal :    Convert a .docx
# History:
# 18/dec/2017    ERK Created
# ----------------------------------------------------------------------------------

import util
import re
from docx import Document

def academic2phonemic(options):
    """COnvert the file in options"""

    # Validate
    if not ('input' in options and 'output' in options):
        return False
    # Make sure we have an error object
    oErr = options['oerr']

    # The maximum characters that we will allow per row
    max_chars_per_row = 80

    try:
        # Get the options
        lStyles = options['styles']
        sTarget = options['target']
        sConvert = options['convert']
        # One of the conversion options is 'interlinear'
        bInterlinear = (sConvert == "interlinear" or sConvert == "i")
        # Get the input and the output file names
        sInput = options['input']
        sOutput = options['output']

        # Create a document object
        doc = Document(sInput)
        # Get the styles in the document
        styles = doc.styles
        # Get target style
        target = next((item for item in styles if item.name == sTarget), None )
        
        iPar = 0
        # Walk through all the paragraphs of the document
        oErr.Status("Walking paragraphs...")
        for par in doc.paragraphs:
            iPar += 1
            # Show where we are
            oErr.Status("  paragraph #{}".format(iPar), True)
            # Check paragraph
            if par.style.name in lStyles:
                oErr.Status("Convert par within document")
                # Convert this part
                t = do_convert(par.text, options)
                # Replace it
                par.text = t
            for run in par.runs:
                s = run.style.name
                if s in lStyles:
                    # Convert this part
                    t = do_convert(run.text, options)
                    # Replace it
                    run.text = t
                    # if needed change style name
                    if target:
                        run.style = target

        # Next visit all TABLES in the document
        iTbl = 0
        table_list = []
        oErr.Status("Walking tables...")
        for tbl in doc.tables:
            iTbl += 1
            # Show where we are
            oErr.Status("  table #{}".format(iTbl), True)

            # Find out how many columns there are


            # Visit all rows
            iRow = 0
            for row in tbl.rows:
                iRow += 1
                iCel = 0
                lstCells = []
                # And then all the cells in this row
                for cell in row.cells:
                    # Check if the cell is already there
                    if not cell in lstCells:
                        lstCells.append(cell)
                        iCel += 1
                        iPar = 0
                        # Check the paragraphs in this cell
                        for par in cell.paragraphs:
                            iPar += 1
                            if par.style.name in lStyles:
                                oErr.Status("Convert par within table-cell")
                                # Convert this part
                                t = do_convert(par.text, options)
                                # Replace it
                                par.text = t
                            for run in par.runs:
                                s = run.style.name
                                if s in lStyles:
                                    # Convert this part
                                    t = do_convert(run.text, options)
                                    # Replace it
                                    run.text = t
                                    # if needed change style name
                                    if target:
                                        run.style = target

            # If we have the interlinearisation option set, then we need to perform that conversion
            if bInterlinear and len(tbl.columns) == 2:
                table_list.append(tbl)

        # Now treat the first table
        if len(table_list) > 0:
            # Take the first table
            tbl = table_list[0]

            # Yes, interlinearisation: Walk the source
            num_gloss = 0       # Characters in the gloss-line
            num_morph = 0       # Characters in the morph-line
            tblNew = None
            for row in tbl.rows:
                # Get the cells
                cell_gloss = row.cells[0]
                cell_morph = row.cells[1]
                # Get text content
                txt_gloss = cell_gloss.text
                txt_morph = cell_morph.text
                # Check counts
                if num_gloss + len(txt_gloss) > max_chars_per_row or num_morph + len(txt_morph) > max_chars_per_row or tblNew == None:
                    # Create a new table with two rows and zero columns
                    tblNew = doc.add_table(2,0)
                    tblNew.autofit = True
                    oErr.Status("Added table. THe total tables = {}".format(len(doc.tables)))
                    num_gloss = 0       # Characters in the gloss-line
                    num_morph = 0       # Characters in the morph-line
                # Add a column to the existing [tblNew]
                colNew = tblNew.add_column(10)
                # Copy cells
                copy_table_cell(oErr, cell_gloss, colNew.cells[0])
                copy_table_cell(oErr, cell_morph, colNew.cells[1])
                # Adapt sizes
                num_gloss += len(txt_gloss)
                num_morph += len(txt_morph)

            # Remove the original table
            # doc.tables.remove(tbl)


        # Save the document under the new name
        doc.save(sOutput)

        # Return okay
        return True
    except:
        oErr.DoError("academic2phonemic")
        return False

def copy_table_cell(oErr, cell_src, cell_dst):
    """Copy the contents of a table cell from source to destination"""

    try:
        # Check the paragraphs in this cell
        for par in cell_src.paragraphs:
            # Add a paragraph to the destination
            par_dst = cell_dst.add_paragraph(text=par.text)
            # Copy the runs of this paragraph
            for run in par.runs:
                s = run.style.name
                run_dst = par_dst.add_run(run.text)
                run_dst.style = run.style
        # Return positively
        return True
    except:
        oErr.DoError("copy_table_cell")
        return False


def do_convert(sPart, options=None):
    """Convert the string in [sPart] according to the rules for the Caucasian handbook"""

    # Get the list of switches
    switches = [] if not 'switches' in options else options['switches']
    bGeminateV = ('vv' in switches)
    bGeminateC = ('cc' in switches)
    bIngush = ('ingush' in switches)
    bHw = ('hw' in switches)
    bGh = ('gh' in switches)
    bOld = False

    # Treat the 'w' where it is a hw occurring after: c, ch, k, p, sh, s, t
    sPart = re.sub(r"(ch|c|k|p|sh|s|t)w", r"\g<1>ħ",sPart)

    if bGeminateC:
        # Treat 'ww'
        sPart = sPart.replace("ww", "ʕʕ")
        if not bGh:
            # Convert gh > ʁ (long and short variant)
            sPart = sPart.replace("ggh", "ʁː").replace("gh", "ʁ").replace("Gh", "ʁ")
        # Convert ch > č (long and short variant)
        sPart = sPart.replace("cch", "čč").replace("ch", "č").replace("Ch", "č")
        # Convert zh > ž (long and short variant)
        sPart = sPart.replace("zzh", "žž").replace("zh", "ž").replace("Zh", "ž")
        # Convert sh > š (long and short variant)
        sPart = sPart.replace("ssh", "šš").replace("sh", "š").replace("Sh", "š")
        if not bHw:
            # Convert hw > ħ (long and short variant)
            sPart = sPart.replace("hhw", "ħħ").replace("hw", "ħ").replace("Hw", "ħ")
    else:
        # Treat 'ww'
        sPart = sPart.replace("ww", "ʕː")
        if not bGh:
            # Convert gh > ʁ (long and short variant)
            sPart = sPart.replace("ggh", "ʁː").replace("gh", "ʁ").replace("Gh", "ʁ")
        # Convert ch > č (long and short variant)
        sPart = sPart.replace("cch", "čː").replace("ch", "č").replace("Ch", "č")
        # Convert zh > ž (long and short variant)
        sPart = sPart.replace("zzh", "žː").replace("zh", "ž").replace("Zh", "ž")
        # Convert sh > š (long and short variant)
        sPart = sPart.replace("ssh", "šː").replace("sh", "š").replace("Sh", "š")
        if not bHw:
            # Convert hw > ħ (long and short variant)
            sPart = sPart.replace("hhw", "ħː").replace("hw", "ħ").replace("Hw", "ħ")
    # Treat the 'w' where it occurs in other places
    sPart = re.sub(r"[Ww]", r"ʕ",sPart)
    # Treat double glottal stop
    if bGeminateC:
        sPart = sPart.replace("''", "ʔʔ")
        sPart = sPart.replace("’’", "ʔʔ")
    else:
        sPart = sPart.replace("''", "ʔː")
        sPart = sPart.replace("’’", "ʔː")
    # Treat single glottal stop
    sPart = re.sub(r"([aeiuoy])(['’])", r"\g<1>ʔ", sPart)
    # Treat [rh]
    sPart = sPart.replace("rh", "r̥")

    # The [v] does NOT change!!!
    
    # Make sure that ejectives have the correct apostrophe
    sPart = sPart.replace("'", "’")
    if bGeminateV:
        if not bIngush:
            sPart = sPart.replace("Yy", "üː").replace("yy", "üː")
    else:
        # Long vowels
        sPart = sPart.replace("Aa", "aː").replace("aa", "aː")
        sPart = sPart.replace("Ee", "eː").replace("ee", "eː")
        sPart = sPart.replace("Ii", "iː").replace("ii", "iː")
        sPart = sPart.replace("Oo", "oː").replace("oo", "oː")
        sPart = sPart.replace("Uu", "uː").replace("uu", "uː")
        if not bIngush:
            sPart = sPart.replace("Yy", "üː").replace("yy", "üː")
    # Diphthong
    if bOld:
        sPart = sPart.replace("Ye", "üe").replace("ye", "üe")
        sPart = sPart.replace("Oe", "üe").replace("oe", "üe")   # So /ye/ and /oe/ coincide
        sPart = sPart.replace("Ov", "ou").replace("ov", "ou")
        sPart = sPart.replace("Ev", "eü").replace("ev", "eü")
        sPart = sPart.replace("Av", "au").replace("av", "au")
    else:
        sPart = re.sub(r"Ye(?:^[aeiouy])", r"üe", sPart)
        sPart = re.sub(r"Oe(?:^[aeiouy])", r"üe", sPart)
        sPart = re.sub(r"Ov(?:^[aeiouy])", r"ou", sPart)
        sPart = re.sub(r"Ev(?:^[aeiouy])", r"eü", sPart)
        sPart = re.sub(r"Av(?:^[aeiouy])", r"au", sPart)
        sPart = re.sub(r"ye(?:^[aeiouy])", r"üe", sPart)
        sPart = re.sub(r"oe(?:^[aeiouy])", r"üe", sPart)
        sPart = re.sub(r"ov(?:^[aeiouy])", r"ou", sPart)
        sPart = re.sub(r"ev(?:^[aeiouy])", r"eü", sPart)
        sPart = re.sub(r"av(?:^[aeiouy])", r"au", sPart)
        pass
    # SHort vowels
    if not bIngush:
        sPart = sPart.replace("Y", "ü").replace("y", "ü")
    # Long consonants
    if not bGeminateC:
        sPart = sPart.replace("bb", "bː").replace("dd", "dː").replace("gg", "gː")
        sPart = sPart.replace("hh", "hː").replace("kk", "kː").replace("ll", "lː")
        sPart = sPart.replace("mm", "mː").replace("nn", "nː").replace("pp", "pː")
        sPart = sPart.replace("qq", "qː").replace("rr", "rː").replace("ss", "sː")
        sPart = sPart.replace("tt", "tː").replace("vv", "vː").replace("xx", "xː")
        sPart = sPart.replace("zz", "zː")

    # Return what we have made of it
    return sPart
    
